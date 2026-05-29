"""
PDFtoDocHub - Free PDF & Word Converter
No API. No login. No cost. Pure Python.
"""

import os, uuid, shutil, subprocess, threading, time
from flask import Flask, request, jsonify, send_file, render_template, abort
import pdfplumber
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

@app.route('/google1234abcd.html')
def google_verify():
    return send_file('static/google1234abcd.html')
    
app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024  # 50 MB limit

UPLOAD_FOLDER    = "uploads"
CONVERTED_FOLDER = "converted"
ALLOWED_EXT      = {"pdf", "docx", "doc"}

os.makedirs(UPLOAD_FOLDER,    exist_ok=True)
os.makedirs(CONVERTED_FOLDER, exist_ok=True)


# ── helpers ──────────────────────────────────────────────────────────────────

def allowed(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXT

def cleanup_file(path, delay=300):
    """Delete a file after `delay` seconds (default 5 min)."""
    def _del():
        time.sleep(delay)
        try:
            os.remove(path)
        except FileNotFoundError:
            pass
    threading.Thread(target=_del, daemon=True).start()


# ── conversion logic ─────────────────────────────────────────────────────────

def pdf_to_word(pdf_path, out_path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    with pdfplumber.open(pdf_path) as pdf:
        total = len(pdf.pages)
        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text()
            if not text or not text.strip():
                doc.add_paragraph(f"[Page {i} – no extractable text]").italic = True
            else:
                for line in text.splitlines():
                    line = line.strip()
                    if not line:
                        doc.add_paragraph()
                        continue
                    para = doc.add_paragraph()
                    is_heading = len(line) < 80 and line.isupper() and not line.endswith(".")
                    run = para.add_run(line)
                    run.bold      = is_heading
                    run.font.size = Pt(13 if is_heading else 11)

            # tables
            for tbl_data in (page.extract_tables() or []):
                if not tbl_data:
                    continue
                rows = len(tbl_data)
                cols = max(len(r) for r in tbl_data)
                tbl  = doc.add_table(rows=rows, cols=cols)
                tbl.style = "Table Grid"
                for ri, row in enumerate(tbl_data):
                    for ci, cell in enumerate(row):
                        c = tbl.cell(ri, ci)
                        c.text = str(cell or "")
                        if ri == 0:
                            c.paragraphs[0].runs[0].bold = True
                doc.add_paragraph()

            if i < total:
                doc.add_page_break()

    doc.save(out_path)


def word_to_pdf(docx_path, out_path):
    # Search all possible LibreOffice locations
    lo = (
        shutil.which("libreoffice") or
        shutil.which("soffice") or
        shutil.which("libreoffice7.6") or
        "/usr/bin/libreoffice" if os.path.exists("/usr/bin/libreoffice") else None or
        "/usr/bin/soffice"     if os.path.exists("/usr/bin/soffice")     else None or
        "/nix/var/nix/profiles/default/bin/libreoffice"
    )

    if not lo or not os.path.exists(lo):
        raise EnvironmentError(
            "LibreOffice not found. Make sure nixpacks.toml has: nixPkgs = [\"libreoffice\"]"
        )

    out_dir = os.path.dirname(os.path.abspath(out_path))
    result  = subprocess.run(
        [lo, "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
        capture_output=True, text=True, timeout=60
    )
    if result.returncode != 0:
        raise RuntimeError(result.stderr or "LibreOffice conversion failed")

    base   = os.path.splitext(os.path.basename(docx_path))[0]
    lo_pdf = os.path.join(out_dir, base + ".pdf")
    if os.path.abspath(lo_pdf) != os.path.abspath(out_path) and os.path.exists(lo_pdf):
        shutil.move(lo_pdf, out_path)


# ── routes ────────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/convert", methods=["POST"])
def convert():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    f    = request.files["file"]
    mode = request.form.get("mode", "")   # "pdf2word" | "word2pdf"

    if not f.filename or not allowed(f.filename):
        return jsonify({"error": "Invalid file type"}), 400

    uid      = uuid.uuid4().hex
    ext      = f.filename.rsplit(".", 1)[1].lower()
    in_path  = os.path.join(UPLOAD_FOLDER, f"{uid}.{ext}")
    f.save(in_path)
    cleanup_file(in_path)

    try:
        if mode == "pdf2word":
            if ext != "pdf":
                return jsonify({"error": "Upload a PDF for PDF→Word"}), 400
            out_name = f"{uid}_converted.docx"
            out_path = os.path.join(CONVERTED_FOLDER, out_name)
            pdf_to_word(in_path, out_path)
            dl_name  = os.path.splitext(f.filename)[0] + ".docx"

        elif mode == "word2pdf":
            if ext not in ("docx", "doc"):
                return jsonify({"error": "Upload a Word file for Word→PDF"}), 400
            out_name = f"{uid}_converted.pdf"
            out_path = os.path.join(CONVERTED_FOLDER, out_name)
            word_to_pdf(in_path, out_path)
            dl_name  = os.path.splitext(f.filename)[0] + ".pdf"

        else:
            return jsonify({"error": "Invalid conversion mode"}), 400

    except Exception as e:
        return jsonify({"error": str(e)}), 500

    cleanup_file(out_path)
    return jsonify({"download_id": out_name, "filename": dl_name})


@app.route("/download/<path:filename>")
def download(filename):
    # prevent path traversal
    safe = os.path.basename(filename)
    path = os.path.join(CONVERTED_FOLDER, safe)
    if not os.path.exists(path):
        abort(404)
    dl_name = request.args.get("as", safe)
    return send_file(path, as_attachment=True, download_name=dl_name)


if __name__ == "__main__":
    app.run(debug=True, port=5000)
